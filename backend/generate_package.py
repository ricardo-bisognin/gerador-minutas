from __future__ import annotations

from pathlib import Path
from typing import Any

from docx_generator import generate_termo_cooperacao
from pt_generator import generate_plano_trabalho


def generate_package(
    tc_template: str | Path,
    pt_template: str | Path,
    output_dir: str | Path,
    data: dict[str, Any],
) -> dict[str, Path]:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    termo = generate_termo_cooperacao(
        tc_template,
        output / "Termo_de_Cooperacao.docx",
        data,
    )
    plano = generate_plano_trabalho(
        pt_template,
        output / "Plano_de_Trabalho.docx",
        data,
    )

    result = {"termo_cooperacao": termo, "plano_trabalho": plano}

    if float(data.get("vigencia_anos", 5)) < 5:
        justificativa = str(data.get("vigencia_justificativa", "")).strip()
        if not justificativa:
            raise ValueError("Vigência inferior a 5 anos exige justificativa.")
        justificativa_path = output / "Justificativa_Vigencia_Inferior_5_Anos.docx"
        _write_justificativa(justificativa_path, data, justificativa)
        result["justificativa_vigencia"] = justificativa_path

    return result


def _write_justificativa(path: Path, data: dict[str, Any], justificativa: str) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("Justificativa para Vigência Inferior a 5 (Cinco) Anos", level=1)
    doc.add_paragraph(
        f"Termo de Cooperação — {data.get('convenente_nome_razao_social', '[CONVENENTE]')}"
    )
    doc.add_paragraph(
        f"Vigência proposta: {data.get('vigencia_anos', '[PREENCHER]')} anos."
    )
    doc.add_paragraph(justificativa)
    doc.add_paragraph(
        "A presente justificativa será juntada ao processo administrativo para análise e decisão da autoridade competente."
    )
    doc.save(path)
