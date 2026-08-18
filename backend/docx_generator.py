from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def _all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def _replace_paragraph_text(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    new_text = text
    for old, new in replacements.items():
        new_text = new_text.replace(old, new)
    if new_text == text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def _format_number(value: int) -> str:
    words = {1:"um",2:"dois",3:"três",4:"quatro",5:"cinco",6:"seis",7:"sete",8:"oito",9:"nove",10:"dez",20:"vinte",30:"trinta",40:"quarenta",50:"cinquenta",60:"sessenta",70:"setenta",80:"oitenta",90:"noventa",100:"cem"}
    return words.get(value, str(value))


def _list_text(values: list[str], empty: str) -> str:
    clean = [str(v).strip() for v in values if str(v).strip()]
    if not clean:
        return empty
    if len(clean) == 1:
        return clean[0]
    if len(clean) == 2:
        return f"{clean[0]} e {clean[1]}"
    return ", ".join(clean[:-1]) + f" e {clean[-1]}"


def _establishments_text(establishments: list[dict[str, Any]]) -> str:
    return _list_text([str(item.get("nome", "")) for item in establishments], "[SELECIONAR ESTABELECIMENTO]")


def _activities_text(activities: list[str]) -> str:
    return _list_text(activities, "[INFORMAR ATIVIDADES]")


def build_replacements(data: dict[str, Any]) -> dict[str, str]:
    quantidade = int(data.get("quantidade_pessoas", 0))
    vigencia = float(data.get("vigencia_anos", 5))
    inicio = int(data.get("periodo_inicio_ano", data.get("ano_assinatura", 2026)))
    fim = int(data.get("periodo_fim_ano", inicio + int(vigencia)))
    vigencia_texto = "5 (cinco) anos" if vigencia == 5 else f"{vigencia:g} anos"
    return {
        "{{fpe_formatado}}": str(data.get("fpe_formatado") or "[PREENCHER FPE]"),
        "{{processo_formatado}}": str(data.get("processo_formatado") or "[PREENCHER PROCESSO]"),
        "{{convenente_nome_razao_social}}": str(data.get("convenente_nome_razao_social", "[INFORMAR CONVENENTE]")),
        "{{convenente_endereco_completo}}": str(data.get("convenente_endereco_completo", "[INFORMAR ENDEREÇO]")),
        "{{convenente_endereco}}": str(data.get("convenente_endereco", data.get("convenente_endereco_completo", "[INFORMAR ENDEREÇO]"))),
        "{{convenente_cnpj}}": str(data.get("convenente_cnpj", "[INFORMAR CNPJ]")),
        "{{representante_nome}}": str(data.get("representante_nome", "[INFORMAR REPRESENTANTE]")),
        "{{representante_rg}}": str(data.get("representante_rg", "[INFORMAR RG]")),
        "{{representante_cpf}}": str(data.get("representante_cpf", "[INFORMAR CPF]")),
        "{{representante_cargo}}": str(data.get("representante_cargo", "[INFORMAR CARGO]")),
        "{{jornada_horario}}": str(data.get("jornada_horario", "[INFORMAR HORÁRIO]")),
        "{{jornada_intervalo}}": str(data.get("jornada_intervalo", "[INFORMAR INTERVALO]")),
        "{{jornada_dias}}": str(data.get("jornada_dias", "[INFORMAR DIAS]")),
        "{{atividades_formatadas}}": _activities_text(data.get("atividades", [])),
        "{{local_prestacao}}": str(data.get("local_prestacao", "[INFORMAR LOCAL DE PRESTAÇÃO]")),
        "{{quantidade_pessoas_formatada}}": f"{quantidade} ({_format_number(quantidade)})" if quantidade else "[INFORMAR QUANTIDADE]",
        "{{quantidade_pessoas}}": str(quantidade) if quantidade else "[INFORMAR QUANTIDADE]",
        "{{regimes_formatados}}": str(data.get("regimes_formatados", "[INFORMAR REGIMES]")),
        "{{estabelecimento_nome}}": _establishments_text(data.get("estabelecimentos", [])),
        "{{estabelecimentos_formatados}}": _establishments_text(data.get("estabelecimentos", [])),
        "{{remuneracao_texto}}": str(data.get("remuneracao_texto", "[INFORMAR REMUNERAÇÃO]")),
        "{{remuneracao_plano_aplicacao}}": str(data.get("remuneracao_texto", "[INFORMAR REMUNERAÇÃO]")),
        "{{vigencia_formatada}}": vigencia_texto,
        "{{ano_assinatura}}": str(data.get("ano_assinatura", inicio)),
        "{{periodo_inicio_ano}}": str(inicio),
        "{{periodo_fim_ano}}": str(fim),
        "{{quantidade_meses}}": str(int(vigencia * 12)),
        "{{ssps_representante_nome}}": str(data.get("ssps_representante_nome", "[CONFIGURAR REPRESENTANTE SSPS]")),
        "{{ssps_representante_rg}}": str(data.get("ssps_representante_rg", "[CONFIGURAR RG SSPS]")),
        "{{ssps_representante_cpf}}": str(data.get("ssps_representante_cpf", "[CONFIGURAR CPF SSPS]")),
        "{{ssps_matricula}}": str(data.get("ssps_matricula", "[CONFIGURAR MATRÍCULA SSPS]")),
        "{{pp_representante_nome}}": str(data.get("pp_representante_nome", "[CONFIGURAR REPRESENTANTE POLÍCIA PENAL]")),
        "{{pp_representante_rg}}": str(data.get("pp_representante_rg", "[CONFIGURAR RG POLÍCIA PENAL]")),
        "{{pp_representante_cpf}}": str(data.get("pp_representante_cpf", "[CONFIGURAR CPF POLÍCIA PENAL]")),
        "{{pp_matricula}}": str(data.get("pp_matricula", "[CONFIGURAR MATRÍCULA POLÍCIA PENAL]")),
    }


def generate_termo_cooperacao(template_path: str | Path, output_path: str | Path, data: dict[str, Any]) -> Path:
    document = Document(str(template_path))
    replacements = build_replacements(data)
    for paragraph in _all_paragraphs(document):
        _replace_paragraph_text(paragraph, replacements)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
