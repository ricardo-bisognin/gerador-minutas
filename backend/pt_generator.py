from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from jornada import criar_jornada


def _replace_all(document: Document, replacements: dict[str, str]) -> None:
    paragraphs = list(document.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)

    for paragraph in paragraphs:
        old = paragraph.text
        new = old
        for token, value in replacements.items():
            new = new.replace(token, value)
        if new != old:
            if paragraph.runs:
                paragraph.runs[0].text = new
                for run in paragraph.runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(new)


def _list(values: list[str]) -> str:
    values = [str(v).strip() for v in values if str(v).strip()]
    if not values:
        return "[PREENCHER]"
    if len(values) == 1:
        return values[0]
    if len(values) == 2:
        return f"{values[0]} e {values[1]}"
    return ", ".join(values[:-1]) + f" e {values[-1]}"


def build_pt_replacements(data: dict[str, Any]) -> dict[str, str]:
    vigencia = float(data.get("vigencia_anos", 5))
    inicio = data.get("periodo_execucao_inicio", "[PREENCHER INÍCIO]")
    fim = data.get("periodo_execucao_fim", "[PREENCHER TÉRMINO]")
    estabelecimentos = data.get("estabelecimentos", [])
    jornadas = criar_jornada(
        data.get("jornada_dias_trabalhados", ["segunda-feira", "terça-feira", "quarta-feira", "quinta-feira", "sexta-feira"]),
        float(data.get("jornada_horas_semanais", 40)),
        str(data.get("jornada_horario", "")),
        str(data.get("jornada_intervalo", "")),
    )

    replacements = {
        "{{convenente_cnpj}}": str(data.get("convenente_cnpj", "[PREENCHER CNPJ]")),
        "{{convenente_nome_razao_social}}": str(data.get("convenente_nome_razao_social", "[PREENCHER RAZÃO SOCIAL]")),
        "{{convenente_endereco}}": str(data.get("convenente_endereco", data.get("convenente_endereco_completo", "[PREENCHER ENDEREÇO]"))),
        "{{convenente_endereco_completo}}": str(data.get("convenente_endereco_completo", "[PREENCHER ENDEREÇO]")),
        "{{convenente_municipio}}": str(data.get("convenente_municipio", "[PREENCHER MUNICÍPIO]")),
        "{{convenente_uf}}": str(data.get("convenente_uf", "RS")),
        "{{convenente_cep}}": str(data.get("convenente_cep", "[PREENCHER CEP]")),
        "{{convenente_telefone}}": str(data.get("convenente_telefone", "[PREENCHER TELEFONE]")),
        "{{convenente_email}}": str(data.get("convenente_email", "[PREENCHER E-MAIL]")),
        "{{convenente_homepage}}": str(data.get("convenente_homepage", "[PREENCHER HOME PAGE]")),
        "{{representante_nome}}": str(data.get("representante_nome", "[PREENCHER REPRESENTANTE]")),
        "{{representante_cpf}}": str(data.get("representante_cpf", "[PREENCHER CPF]")),
        "{{representante_rg}}": str(data.get("representante_rg", "[PREENCHER RG]")),
        "{{representante_orgao_expedidor}}": str(data.get("representante_orgao_expedidor", "[PREENCHER ÓRGÃO EXPEDIDOR]")),
        "{{representante_cargo}}": str(data.get("representante_cargo", "[PREENCHER CARGO]")),
        "{{representante_funcao}}": str(data.get("representante_funcao", "[PREENCHER FUNÇÃO]")),
        "{{periodo_execucao_inicio}}": str(inicio),
        "{{periodo_execucao_fim}}": str(fim),
        "{{vigencia_meses}}": str(int(vigencia * 12)),
        "{{atividades}}": _list(data.get("atividades", [])),
        "{{estabelecimentos}}": _list([str(e.get("nome", "")) for e in estabelecimentos]),
        "{{jornada_dias_trabalhados}}": jornadas.dias_trabalhados_formatados,
        "{{jornada_dias_descanso}}": jornadas.dias_descanso_formatados,
        "{{jornada_horas_semanais}}": str(data.get("jornada_horas_semanais", "[PREENCHER]")),
        "{{jornada_horario}}": str(data.get("jornada_horario", "[PREENCHER]")),
        "{{jornada_intervalo}}": str(data.get("jornada_intervalo", "[PREENCHER]")),
        "{{remuneracao_texto}}": str(data.get("remuneracao_texto", "[PREENCHER REMUNERAÇÃO]")),
        "{{quantidade_pessoas}}": str(data.get("quantidade_pessoas", "[PREENCHER QUANTIDADE]")),
    }
    return replacements


def generate_plano_trabalho(template_path: str | Path, output_path: str | Path, data: dict[str, Any]) -> Path:
    document = Document(str(template_path))
    _replace_all(document, build_pt_replacements(data))
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
